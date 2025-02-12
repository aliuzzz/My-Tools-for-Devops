########
#基于cacti 1.2.27+版本
import requests,logging,time,calendar,os,shutil,cairosvg
import concurrent.futures
import matplotlib.pyplot as plt
from requests.adapters import HTTPAdapter
from datetime import datetime
from dateutil.relativedelta import relativedelta
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,Cm,RGBColor
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

chrome_options = Options()
chrome_options.add_argument('--headless') 
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-dev-shm-usage')
chrome_options.add_argument('--window-size=1920,1080')
driver = webdriver.Chrome(options=chrome_options)
driver.get('http://10.105.1.253/cacti/')
user = 'xxxxx'
passwd = 'xxxxxxxxxx'

# 读当前文件夹下的csv文件，并将数据写入字典
df = pd.read_csv('D://Users//Worker//Desktop//test//monitor95//monthly//get_info.csv')
num_list = {}
for i in range(len(df)):
    num_list[df.iloc[i][0]] = [df.iloc[i][1],df.iloc[i][2],df.iloc[i][3]]
groups = df.groupby('company')['description'].apply(list)


#----------------处理“443值”跟月份日期的问题------------------------------
today = datetime.now()
today_year = today.year
today_month = today.month
previous_month = today - relativedelta(months=1)
#计算“443”值的公式,28天--404 30天-433 31天--447
month_range = calendar.monthrange(previous_month.year, previous_month.month)
num_days = month_range[1]
formula_443 = round(288 * num_days * 0.05) + 1
#----------------------------------------------------------------------

# 清空文件夹,并创建全新的文件夹
def clear_file():
    info_path = "D://Users//Worker//Desktop//test//monitor95//monthly//info"
    for root, dirs, files in os.walk(info_path):
        for file in files:
            os.remove(os.path.join(root, file))
        for dir in dirs:
            shutil.rmtree(os.path.join(root, dir))
    time.sleep(2)
    for i in range(len(df)):
        dirs_name = str(df.values[i,1])
        path_charts = "D://Users//Worker//Desktop//test//monitor95//monthly//info//charts/" + dirs_name +"/"
        path_csv = "D://Users//Worker//Desktop//test//monitor95//monthly//info//csv/" + dirs_name +"/"      
        path_prtsc = "D://Users//Worker//Desktop//test//monitor95//monthly//info//prtscreen/" + dirs_name + "/"
        os.makedirs(path_charts)
        os.makedirs(path_csv)
        os.makedirs(path_prtsc)
    os.makedirs("D://Users//Worker//Desktop//test//monitor95//monthly//info//to_PrtS/")
    logging.info('文件夹创建成功')

#登陆下载
def selenium_login():
    # ----------selenium登录------------------------------------------------
    driver.maximize_window() 
    login_user = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, '[name="login_username"]')))
    login_password = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, '[name="login_password"]')))
    login_user.send_keys(user)
    login_password.send_keys(passwd)
    driver.find_element(By.CSS_SELECTOR,'[value="登录"]').click()
    # 保存 session
    session = requests.Session()
    # 创建HTTPAdapter对象，并设置连接池大小
    adapter = HTTPAdapter(pool_maxsize=100)

    # 将HTTPAdapter对象传递给Session对象的mount方法
    session.mount('http://', adapter)
    session.mount('https://', adapter)
    for cookie in driver.get_cookies():
        if cookie['name'] == 'Cacti':
            print(f"Name: {cookie['name']}")
            print(f"Value: {cookie['value']}")
    session.cookies.set(cookie['name'], cookie['value'])

    headers = {
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/99.0.4844.51 Safari/537.36',
        'content-type': 'text/html; charset=utf-8',
        'Referer': 'http://10.105.1.253/cacti/graph_view.php'
    }
    # 使用 session 发送请求
    response = session.get('http://10.105.1.253/cacti/graph_view.php', headers=headers, stream=True)
    print(response)
    # ------------------------下载文件----------------------------------------
    def download_file(num, value, start_time, end_time, file_type):
        if file_type == 'pic':
            download_url = f'http://10.105.1.253/cacti/graph_image.php?local_graph_id={num}&rra_id=0&graph_height=250&graph_width=700&graph_start={start_time}&graph_end={end_time}'
            print(download_url)
            file_dir = f'D://Users//Worker//Desktop//test//monitor95//monthly//info//charts/{value[0]}/'
            file_name = f'{value[0]}.svg'
        else:
            download_url = f'http://10.105.1.253/cacti/graph_xport.php?local_graph_id={num}&rra_id=0&view_type=tree&graph_start={start_time}&graph_end={end_time}'
            print(download_url)
            file_dir = f'D://Users//Worker//Desktop//test//monitor95//monthly//info//csv//{value[0]}/'
            file_name = f'{value[0]}.csv'
        file_path = file_dir + file_name
        if not os.path.exists(file_dir):
            os.makedirs(file_dir)
        with open(file_path, "wb") as f:
            r = session.get(download_url, headers=headers, stream=True)
            total_length = r.headers.get('content-length')
            if total_length is None: 
                f.write(r.content)
            else:
                dl = 0
                total_length = int(total_length)
                for data in r.iter_content(chunk_size=4096):
                    dl += len(data)
                    f.write(data)
                # 判断文件大小是否为0，如果是，则重新下载一次
                if dl == 0:
                    r = session.get(download_url, headers=headers, stream=True)
                    for data in r.iter_content(chunk_size=4096):
                        f.write(data)
        logging.info(f'Download {file_type} success: {file_path}')
    # --------------------------处理日期------------------------------------------
    thismonth_first_day = datetime(today_year, today_month, 1)
    thismonth_zero_time = int(thismonth_first_day.timestamp())
    lastmonth_first_day = today - relativedelta(months=1)
    lastmonth_first_day = datetime(lastmonth_first_day.year, lastmonth_first_day.month, 1)
    lastmonth_zero_time = int(lastmonth_first_day.timestamp())
    # --------------------------下载文件------------------------------------
    with concurrent.futures.ThreadPoolExecutor(max_workers=30) as executor: 
        print('开始下载文件')
        for num, value in num_list.items():
            executor.submit(download_file, num, value, lastmonth_zero_time, thismonth_zero_time, 'pic') 
            executor.submit(download_file, num, value, lastmonth_zero_time, thismonth_zero_time, 'csv')
    driver.close()
    driver.quit()

def csv_deal():
    # 读文件-------------------------------------------------------------------
    directory_name = r'D://Users//Worker//Desktop//test//monitor95//monthly//info//csv/'
    prtsc_name = r'D://Users//Worker//Desktop//test//monitor95//monthly//info//prtscreen/'
    for num,value in num_list.items():
        directory_names = directory_name + str(value[0])
        prtsc_names = prtsc_name +str(value[0])
        filenames = os.listdir(directory_names)
    # ------------------------------------------------------------------------ 
        for csvname in filenames:
            filename = (directory_names+"/"+csvname)
            filename = os.path.join(directory_names, csvname)
    # -----------处理表格-----------------------------------------------------
        print(filename)
        # 找到第一个空行位置
        blank_line = 0
        with open(filename,'r',encoding='utf8') as f:
            for (num,value) in enumerate(f):
                if value.strip() == '""':
                    blank_line = num + 1
                    break
        df = pd.read_csv(filename, skiprows=blank_line)       
        cols_to_keep = df.filter(regex='输入|输出').columns
        # 删除不在cols_to_keep中的列
        df.drop(df.columns.difference(cols_to_keep), axis=1, inplace=True)
        # 主从降序排输入值列
        p_col = ['InBound', 'OutBound']
        df.columns = p_col
        df.sort_values(by=['InBound', 'OutBound'], ascending=False, inplace=True)
        df.reset_index(drop=True, inplace=True)
        df['roll-num']=df.index
        df.iloc[:-1]
        # -----------处理数据-----------------------------------------------------    
        print("正在处理输入相关数据信息.......")
        no_start = round(len(df)*0.05)
        no_end = round(len(df)*0.05)+2
        df_in = df.loc[no_start:no_end,['roll-num','InBound']] #抽出需要的行
        df_in = df_in.astype(float).astype(int)  # 这里是添加用来去掉整数后面的多余.00000小数点位
        fig, ax = plt.subplots(figsize=(9, 3)) #设置图像大小（底色画布大小）
        ax.axis('off')

        # 生成表格并设置样式
        table = ax.table(
            cellText=df_in.values,
            colLabels=df_in.columns,
            cellLoc='center',
            loc='center',
            colColours=['#f0f0f0'] * len(df_in.columns)  # 可选：列标题背景色
        )
        table.auto_set_font_size(False)
        table.set_fontsize(15)  # 调整字体大小
        table.scale(1, 3)     # 调整单元格宽高 (宽度, 高度)
        cells = table.get_celld()
        for i in range(len(df_in.columns)):
            if i == 0:
                width = 0.3  # 第一列宽度
            else:
                width = 0.6  # 第二列宽度

            cells[(0, i)].set_width(width)  # 设置列标题宽度
            for j in range(len(df_in)):
                cells[(j + 1, i)].set_width(width)
                if j == 1:  # 修改这个数字来选择要着色的行
                    cells[(j + 1, i)].set_facecolor('#fffa00')
        prtsc_name_in = prtsc_names+'/'+csvname+'-in.png'
        # 保存高清图片
        plt.subplots_adjust(left=0.1, right=0.5, top=0.5, bottom=0.1)
        plt.savefig(prtsc_name_in, dpi=130, bbox_inches='tight',pad_inches=0)
        # 更改顺序，主从降序排输出值列
        df.sort_values(by=['OutBound', 'InBound'], ascending=False, inplace=True)
        df.reset_index(drop=True, inplace=True)
        df['roll-num'] = df.index
        df.iloc[:-1]
        print("正在处理输出相关数据信息.......")
        no_start = round(len(df)*0.05)
        no_end = round(len(df)*0.05)+2
        df_out = df.loc[no_start:no_end,['roll-num','OutBound']]
        df_out = df_out.astype(float).astype(int)
        fig, ax = plt.subplots(figsize=(9, 3)) #设置图像大小（底色画布大小）
        ax.axis('off')

        # 生成表格并设置样式
        table = ax.table(
            cellText=df_in.values,
            colLabels=df_in.columns,
            cellLoc='center',
            loc='center',
            colColours=['#f0f0f0'] * len(df_in.columns)  # 可选：列标题背景色
        )
        table.auto_set_font_size(False)
        table.set_fontsize(15)  # 调整字体大小
        table.scale(1, 3)     # 调整单元格宽高 (宽度, 高度)
        cells = table.get_celld()
        for i in range(len(df_in.columns)):
            if i == 0:
                width = 0.3  # 第一列宽度
            else:
                width = 0.6  # 第二列宽度

            cells[(0, i)].set_width(width)  # 设置列标题宽度
            for j in range(len(df_in)):
                cells[(j + 1, i)].set_width(width)
                if j == 1:  # 修改这个数字来选择要着色的行
                    cells[(j + 1, i)].set_facecolor('#fffa00')
        prtsc_name_out = prtsc_names+'/'+csvname+'-out.png'
        # 保存高清图片
        plt.subplots_adjust(left=0.1, right=0.5, top=0.5, bottom=0.1)
        plt.savefig(prtsc_name_out, dpi=130, bbox_inches='tight',pad_inches=0)
        plt.close()

# 图片拼接
def join_pics():
    for num, value in num_list.items():
        csvprtdir_path = os.path.join('D://Users//Worker//Desktop//test//monitor95//monthly//info//prtscreen/', str(value[0]))
        charts_path = os.path.join('D://Users//Worker//Desktop//test//monitor95//monthly//info//charts//', str(value[0]), f"{value[0]}.svg")
        pic_path = os.listdir(csvprtdir_path)
        
        # 创建合成图片的基础
        csv_result = Image.new('RGB', [1000, 400])  # 表格的最终整体尺寸
        prt_result = Image.new('RGB', [1000, 840], color='white')  # 最终图的整体尺寸
        
        # 处理每一张截图，缩放后添加到列表
        pics = []
        for pic_name in pic_path:
            pic_full_path = os.path.join(csvprtdir_path, pic_name)
            with Image.open(pic_full_path) as img:
                pics.append(img.resize((500, 380)))
        
        # 处理 SVG 图像
        temp_svg_path = 'temp_chart.png'
        cairosvg.svg2png(url=charts_path, write_to=temp_svg_path)
        
        # 打开转换后的 PNG 图像
        charts = Image.open(temp_svg_path).resize((975, 430))

        # 拼接图片
        if len(pics) >= 2:
            csv_result.paste(im=pics[0], box=(0, 0))
            csv_result.paste(im=pics[1], box=(500, 0))  # 右边拼到左边的位置

        prt_result.paste(im=charts, box=(10, 15))  # 流量图的起始位置
        prt_result.paste(im=csv_result, box=(0, 455))  # 表格整体接到上面流量的起始位置
        output_path = os.path.join('D://Users//Worker//Desktop//test//monitor95//monthly//info//to_PrtS', f"{value[0]}.png")
        prt_result.save(output_path)
        os.remove(temp_svg_path)

# 创建文档
def to_docs():
    doc = Document()
    title = str(previous_month.year)+'年'+str(previous_month.month)+'月 对账文档'
    doc.add_heading(title,0)
    style = doc.styles.add_style('DateStyle', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(12)
    style.font.italic = True
    section = doc.sections[0]
    section.top_margin = Cm(0.5) # 设置页边距，单位是Cm
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    make_doc_time = '文档自动化生成时间：'+str(today)
    doc.add_paragraph(make_doc_time, style='DateStyle')
    doc.paragraphs[-1].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #设置正文部分
    for company,description in groups.items():
        print(company)
        t1_1 = doc.add_heading('', level=1)
        t1_1.paragraph_format.space_before = Pt(5)
        t1_1.paragraph_format.space_after = Pt(3)
        t1_1.paragraph_format.line_spacing = Pt(13)
        run_company = t1_1.add_run(company)
        run_company.font.name = 'SimSun'
        run_company.font.bold = True
        run_company.font.size = Pt(14)
        run_company.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        print(description)
        i = 1
        for j in description:
            par1 = doc.add_paragraph(f'({i}) 截图为：{j}')
            par1.paragraph_format.space_before = Pt(0)
            par1.paragraph_format.space_after = Pt(0)
            par1.paragraph_format.line_spacing = Pt(11)
            doc.add_picture('D://Users//Worker//Desktop//test//monitor95//monthly//info//to_PrtS/'+j+'.png',width=Cm(14.5),height=Cm(11.2))
            i +=1
    paragraph = doc.add_paragraph(str(today_year)+'年'+str(today_month)+'月 需要新增/修改部分如下：')
    run = paragraph.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(255,0,0)
    run.font.size = Pt(16)
    doc.save(title+'.docx')

def close_program():
    os._exit(0)

if __name__ == '__main__':
    T1 = time.perf_counter()
    clear_file()
    print('正在下载处理中.....')
    selenium_login()
    csv_deal()
    join_pics()
    to_docs()
    print('完成')
    T2 = time.perf_counter()
    print('程序运行时间:%s秒' % ((T2 - T1)))
    close_program()