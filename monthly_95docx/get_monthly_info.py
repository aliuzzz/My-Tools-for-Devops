import requests,zipfile,logging,datetime, time, calendar,os, shutil
import concurrent.futures
import imgkit  #处理html转图
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.shared import Pt,Cm,RGBColor
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import mimetypes
import smtplib   # 发送邮件模块
from email import encoders
from email.mime.multipart import MIMEMultipart    # 使用MIMEMultipart来标示这个邮件是多个部分组成的
from email.mime.base import MIMEBase
from email.mime.text import MIMEText   # 定义邮件内容
from email.utils import formataddr

chrome_options = Options()
chrome_options.add_argument('--headless') 
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-dev-shm-usage')
chrome_options.add_argument('--window-size=1920,1080')
driver = webdriver.Chrome(options=chrome_options)
driver.get('http://ip/cacti') # cacti所在的ip地址

user = 'xxx' # 用户名    
passwd = 'xxxxxxx' # 密码

# 读当前文件夹下的csv文件，并将数据写入字典
df = pd.read_csv('./get_info.csv')
num_list = {}
for i in range(len(df)):
    num_list[df.iloc[i][0]] = [df.iloc[i][1],df.iloc[i][2],df.iloc[i][3]]
groups = df.groupby('company')['description'].apply(list) # 建一个group方便后面往docx里写的时候分组

# 创建一个log文件记录脚本运行情况
logging.basicConfig(filename='/monthly/monthly_get.log',level=logging.DEBUG,format='%(asctfime)s - %(filename)s[line:%(lineno)d] - %(levelname)s: %(message)s')

#----------------处理“443值”跟月份日期的问题------------------------------
# 当天所在的年份
today_year = datetime.datetime.now().year
# 上个月
today_month = datetime.datetime.now().month-1
# 获取当前月份的总天数,输出的是'(当月第一天是周几,当月有多少天)'
monthRange = calendar.monthrange(today_year,today_month) 
#计算“443”值的公式
formula_443 = round(288*monthRange[1]*0.05)+1 #28天--404 30天-433 31天--447  (443就是95值的位置)
#----------------------------------------------------------------------

# 清空文件夹,并创建全新的文件夹,放在最开头执行，否则该月和上个月的客户不一致时还会多余出来。这里简单粗暴的全清了重下
def clear_file():
    info_path = "/monthly/info"
    for root, dirs, files in os.walk(info_path):
        for file in files:
            os.remove(os.path.join(root, file))
        for dir in dirs:
            shutil.rmtree(os.path.join(root, dir))
    time.sleep(2)
    for i in range(len(df)):
        dirs_name = str(df.values[i,1])
        path_charts = "/monthly/info/charts/" + dirs_name +"/"
        path_csv = "/monthly/info/csv/" + dirs_name +"/"      
        path_prtsc = "/monthly/info/prtscreen/" + dirs_name + "/"
        os.makedirs(path_charts)
        os.makedirs(path_csv)
        os.makedirs(path_prtsc)
    os.makedirs("/monthly/info/html/")
    os.makedirs("/monthly/info/to_PrtS/")
    logging.info('文件夹创建成功')

#登陆下载
def selenium_login():
    # ----------selenium登录------------------------------------------------
    driver.maximize_window() 
    login_user = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, '[name="login_username"]')))
    login_password = WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, '[name="login_password"]')))
    login_user.send_keys(user)
    login_password.send_keys(passwd)
    driver.find_element(By.CSS_SELECTOR,'[value="Login"]').click()
    WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR,'[alt="Graphs"]'))).click()

    # 保存 session
    session = requests.Session()
    for cookie in driver.get_cookies():
        print(cookie)
        session.cookies.set(cookie['name'], cookie['value'])

    headers = {
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/99.0.4844.51 Safari/537.36',
        'content-type': 'text/html; charset=UTF-8',
        'Referer': 'http://ip/cacti/graph_view.php'
    }
    # 使用 session 发送请求
    response = session.get('http://ip/cacti/graph_view.php', headers=headers)

    # ------------------------下载文件----------------------------------------
    def download_file(num, value, start_time, end_time, file_type):
        if file_type == 'pic':# 这里图片的链接上加了宽高参数，可以随即调整，默认是120x500
            download_url = f'http://ip/cacti/graph_image.php?local_graph_id={num}&rra_id=0&graph_height=250&graph_width=700&graph_start={start_time}&graph_end={end_time}'
            file_dir = f'/monthly/info/charts/{value[0]}/'
            file_name = f'{value[0]}.png'
        else:
            download_url = f'http://ip/cacti/graph_xport.php?local_graph_id={num}&rra_id=0&view_type=tree&graph_start={start_time}&graph_end={end_time}'
            file_dir = f'/monthly/info/csv/{value[0]}/'
            file_name = f'{value[0]}.csv'
        file_path = file_dir + file_name
        if not os.path.exists(file_dir):
            os.makedirs(file_dir)
        with open(file_path, "wb") as f:
            r = session.get(download_url, headers=headers, stream=True)
            total_length = r.headers.get('content-length')
            if total_length is None:  # no content length header
                f.write(r.content)
            else:
                dl = 0
                total_length = int(total_length)
                for data in r.iter_content(chunk_size=4096):
                    dl += len(data)
                    f.write(data)
        logging.info(f'Download {file_type} success: {file_path}')

    # --------------------------处理日期------------------------------------------
    # 本月1号零点的时间戳
    thismonth_zero_time = int(time.mktime(datetime.date(datetime.date.today().year,datetime.date.today().month,1).timetuple()))
    # 上月1号零点的时间戳
    lastmonth_zero_time = int(time.mktime(datetime.date(datetime.date.today().year,datetime.date.today().month-1,1).timetuple()))
    # ------------------------------------------------------------------------
    # --------------------------下载文件------------------------------------
    '''
    #单线程下载
    for num, value in num_list.items():
        # 下载图片
        download_file(num, value, lastmonth_zero_time, thismonth_zero_time, 'pic')
        # 下载 csv 文件
        download_file(num, value, lastmonth_zero_time, thismonth_zero_time, 'csv')
    '''
    #使用线程池多线程下载
    with concurrent.futures.ThreadPoolExecutor(max_workers=20) as executor: 
        for num, value in num_list.items(): # 下载图片 
            executor.submit(download_file, num, value, lastmonth_zero_time, thismonth_zero_time, 'pic') 
            # 下载 csv 文件 
            executor.submit(download_file, num, value, lastmonth_zero_time, thismonth_zero_time, 'csv')
    driver.close()
    driver.quit()

def highlight_max(series): #遍历frame的每一列(series)依次设置效果
    is_443th = series == series.iloc[3]
    return ['background-color: #FFC000' if value else '' for value in is_443th]

def csv_deal():
    # 读文件-------------------------------------------------------------------
    directory_name = r'/monthly/info/csv/'
    prtsc_name = r'/monthly/info/prtscreen/'
    for num,value in num_list.items():
        directory_names = directory_name + str(value[0])
        prtsc_names = prtsc_name +str(value[0])
        filenames = os.listdir(directory_names)
    # ------------------------------------------------------------------------ 
        for csvname in filenames:
            filename = (directory_names+"/"+csvname)
            filename = os.path.join(directory_names, csvname)
    # -----------处理表格-----------------------------------------------------
        # 找到第一个空行位置
        blank_line = 0
        with open(filename,'r',encoding='utf8') as f:
            for (num,value) in enumerate(f):
                if value.strip() == '""':
                    blank_line = num + 1
                    break
        # 读文件，跳过无效头部
        df = pd.read_csv(filename, skiprows=blank_line)       
        # 删除第一列
        df.drop(df.columns[0], inplace=True, axis=1)
        # 删除匹配 "col" 的列
        df.drop(list(df.filter(regex='col')), axis=1, inplace=True)
        # 主从降序排输入值列
        p_col = ['输入相关数据信息', '输出相关数据信息']
        df.columns = p_col
        df.sort_values(by=['输入相关数据信息', '输出相关数据信息'], ascending=False, inplace=True)
        df.reset_index(drop=True, inplace=True)
        df['roll-num']=df.index
        #--------这里由于index跟实际相差一行，所以添加这两行用以调整
        df['roll-num']=df['roll-num'].shift(-1) 
        df.iloc[:-1]
        df.style.apply(f,axis=0,subset=['roll-num'])
        # -----------处理数据-----------------------------------------------------    
        print("正在处理输入相关数据信息.......")
        no_start = round(len(df)*0.05)-3
        no_end = round(len(df)*0.05)+3
        df_in = df.loc[no_start:no_end,['roll-num','输入相关数据信息']] #抽出需要的行
        df_in = df_in.astype(float).astype(int)  # 这里是添加用来去掉整数后面的多余.00000小数点位
        df_in = df_in.style.set_properties(**{'background-color':'white',"align":"center"}).apply(highlight_max) 
        html_path_in = '/monthly/info/html/'+csvname+'-in.html'        
        df_in.to_html(html_path_in)
        time.sleep(2)
        # 从这里开始是配置html转图片的部分
        path_wkimg = r'//usr//local//bin//wkhtmltoimage'  #这个要注意需要安装wkhtmltoimage这个东西
        cfg = imgkit.config(wkhtmltoimage=path_wkimg)
        options = {
                'crop-w': 220,#需要截图的宽高以及xy轴的位置，这里可以进行调整
                'crop-h': 800,
                'crop-x': 38,
                'crop-y': 5,
                'encoding': 'utf-8',
            }
        outputpath_in = prtsc_names+'/'+csvname+'-in.png'
        imgkit.from_file(html_path_in, output_path=outputpath_in, config=cfg,options=options)
        # 更改顺序，主从降序排输出值列
        df.sort_values(by=['输出相关数据信息', '输入相关数据信息'], ascending=False, inplace=True)
        df.reset_index(drop=True, inplace=True)
        df['roll-num'] = df.index
        df['roll-num'] = df['roll-num'].shift(-1)
        df.iloc[:-1]
        print("正在处理输出相关数据信息.......")
        no_start = round(len(df)*0.05)-3
        no_end = round(len(df)*0.05)+3
        df_out = df.loc[no_start:no_end,['roll-num','输出相关数据信息']]
        df_out = df_out.astype(float).astype(int)
        df_out = df_out.style.set_properties(**{'background-color':'white',"align":"center"}).apply(highlight_max)
        html_path_out = '/monthly/info/html/'+csvname+'-out.html'
        df_out.to_html(html_path_out)
        outputpath_out = prtsc_names+'/'+csvname+'-out.png'
        imgkit.from_file(html_path_out, output_path=outputpath_out, config=cfg,options=options)

# 图片拼接
def join_pics():
    # 读文件
    for num, value in num_list.items():
        csvprtdir_path = os.path.join('/monthly/info/prtscreen/', str(value[0]))
        charts_path = os.path.join('/monthly/info/charts/', str(value[0]), f"{value[0]}.png")
        pic_path = os.listdir(csvprtdir_path)
        
        # 将多张图片合成一张图片
        csv_result = Image.new('RGB', [500 * 2, 400]) #表格的最终整体尺寸
        prt_result = Image.new('RGB', [500 * 2, 420 * 2], color='white') #最终图的整体尺寸
        
        # 处理每一张截图并将缩放后的图片添加进列表
        pic = []
        for pic_name in pic_path:
            pic_path = os.path.join(csvprtdir_path, pic_name)
            with Image.open(pic_path) as img:
                pic.append(img.resize((500, 380)))
        
        charts = Image.open(charts_path).resize((650, 430))
        width, height = charts.size
        charts = charts.resize((int(width * 1.5), height))

        # 拼接图片
        csv_result.paste(im=pic[0], box=(0, 0))
        csv_result.paste(im=pic[1], box=(500, 0)) #右边拼到左边的位置，坐标系是（x,y）
        prt_result.paste(im=charts, box=(10, 15)) #流量图的起始位置
        prt_result.paste(im=csv_result, box=(0, 455))#表格整体接到上面流量的起始位置
        
        # 将合成的图片保存
        output_path = os.path.join('/monthly/info/to_PrtS', f"{value[0]}.png")
        prt_result.save(output_path)

# 创建文档
def to_docs():
    doc = Document()   
    section = doc.sections[0]
    header = section.header
    section.top_margin = Cm(1.54) # 设置页边距，单位是Cm
    section.bottom_margin = Cm(0.75)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.4)
    
    #设置页眉页脚部分
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(1.75)
    header.paragraphs[0].text = '数据表烦请按照紧急程度截取，打包发送'
    header.paragraphs[0].runs[0] = 'SimSun'
    header.paragraphs[0].runs[0].font.size = Pt(18)
    header.paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,255)
    header.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    header.paragraphs[0].runs[0].font.bold = True
    header.paragraphs[0].runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #设置正文部分
    for company,description in groups.items():
        t1_1 = doc.add_heading('', level=1) #设置heading的属性
        t1_1.paragraph_format.space_before = Pt(5)
        t1_1.paragraph_format.space_after = Pt(3)
        t1_1.paragraph_format.line_spacing = Pt(13)
        run_company = t1_1.add_run(company)
        run_company.font.name = 'SimSun'
        run_company.font.bold = True
        run_company.font.size = Pt(14)
        run_company.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        for j in description:
            par1 = doc.add_paragraph('File name:'+ j)
            par1.paragraph_format.space_before = Pt(0)
            par1.paragraph_format.space_after = Pt(0)
            par1.paragraph_format.line_spacing = Pt(11)
            doc.add_picture('/monthly/info/to_PrtS/'+j+'.png',width=Cm(14.5),height=Cm(11.2)) #这里是加到docx里图片的大小，这个大小刚好一个页面放两个图
    doc.save('/monthly/to_zip/月度对账文档.docx')

# 压缩文件,不传参就不设密码，这个可以再main处填
def zip_files(to_zip_dir_name, zip_password=None):
    # 检查文件夹是否存在
    if not os.path.isdir(to_zip_dir_name) :
        logging.error('文件夹不存在')
        return
    with zipfile.ZipFile('月度对账文档.zip', 'w', zipfile.ZIP_DEFLATED) as zipf:
        if zip_password:
            zipf.setpassword(zip_password.encode('utf-8'))
        # 处理to_zip_dir_name文件夹
        for root, dirs, files in os.walk(to_zip_dir_name):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, arcname=file)  # 写入zip文件
    logging.info('已生成压缩文件')

# 发送邮件
def send_email():
    filepath = "/monthly/月度对账文档.zip"     # 要发送的压缩文件路径
    smtp_server = "smtp.163.com"      # 发送邮箱服务器
    username = "xxxxx@163.com"      # 用于发送邮箱的用户账号
    password = "xxxxxxxxxx"      # 密码(即授权码)
    sender = 'xxxxx@163.com'  # 发送者的邮箱，和username相同即可
    receivers = ['xxxxxxx@126.com']   # 接收者的邮箱
    EMAIL_FROM_NAME = '月度自动化对账'   # 自定义发件名称

    time = datetime.datetime.today().strftime("%m-%d %H：%M")
    msg = MIMEMultipart()
    # 邮件正文
    msg.attach(MIMEText("月度对账文档请查看附件",'plain','utf-8'))   # 文本内容换行\r\n
    msg['From'] = formataddr(pair=(EMAIL_FROM_NAME, sender))     # 自定义发件人的名称
    msg['To'] = receivers[0]  # 发送给receivers里的第一个用户
    #msg['To'] = ";".join(receivers)  # 发送给多个好友
    subject = "{}--月度对账文档".format(time)
    msg['Subject'] = subject

    data = open(filepath, 'rb')
    ctype, encoding = mimetypes.guess_type(filepath)
    if ctype is None or encoding is not None:
        ctype = 'application/octet-stream'
    maintype, subtype = ctype.split('/', 1)
    file_msg = MIMEBase(maintype, subtype)
    file_msg.set_payload(data.read())
    data.close()
    encoders.encode_base64(file_msg)  # 把附件编码
    file_msg.add_header('Content-Disposition', 'attachment', filename="月度对账文档.zip")  # 修改邮件头
    msg.attach(file_msg)
    try:
        server = smtplib.SMTP(smtp_server)
        server.login(username,password)
        server.sendmail(sender,receivers,msg.as_string())
        server.quit()
        logging.info('发送成功')
        logging.info('---------------------End---------------------------------')
    except Exception as err:
        logging.warning('发送失败')
        logging.warning(err)

# 添加一个关闭程序的函数
def close_program():
    os._exit(0)

if __name__ == '__main__':
    T1 = time.perf_counter()
    # 待压缩的文件夹路径
    to_zip_dir_name = '/monthly/to_zip/'
    zip_password = 'xxxxxxxxx'
    clear_file()
    print('正在下载处理中.....')
    selenium_login()
    csv_deal()
    join_pics()
    to_docs()
    zip_files(to_zip_dir_name,zip_password)
    send_email()  #文件过大发不过去，最大50M
    print('完成')
    T2 = time.perf_counter()
    print('程序运行时间:%s秒' % ((T2 - T1)))
    close_program()
    